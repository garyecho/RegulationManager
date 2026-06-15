"""
文档正文提取 — 从 PDF / DOCX / DOC 文件中提取纯文本
"""
import logging
import os
import subprocess
import tempfile
from pathlib import Path

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_path: str) -> str:
    """从 PDF 提取文本（PyMuPDF）"""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(file_path)
        text_parts = []
        for page in doc:
            page_text = page.get_text()
            if page_text.strip():
                text_parts.append(page_text.strip())
        doc.close()
        return "\n".join(text_parts).strip()
    except Exception as e:
        logger.warning(f"PDF 文本提取失败 {file_path}: {e}")
        return ""


def extract_text_from_docx(file_path: str) -> str:
    """从 DOCX 提取文本（python-docx）"""
    try:
        from docx import Document
        doc = Document(file_path)
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())
        # 也提取表格中的文字
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text.strip())
        return "\n".join(text_parts).strip()
    except Exception as e:
        logger.warning(f"DOCX 文本提取失败 {file_path}: {e}")
        return ""


def extract_text_from_doc(file_path: str) -> str:
    """从旧版 .doc 文件提取文本（OLE2 格式）"""
    try:
        # 方法1：尝试 antiword（如果安装了的话）
        try:
            result = subprocess.run(
                ["antiword", file_path],
                capture_output=True, text=True, timeout=30
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout.strip()
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass

        # 方法2：尝试 python-docx 尝试打开（某些 .doc 实际是 .docx 格式）
        try:
            from docx import Document
            doc = Document(file_path)
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())
            if text_parts:
                return "\n".join(text_parts).strip()
        except Exception:
            pass

        # 方法3：UTF-16LE 解码提取（Word 内部存储格式）
        try:
            import re
            with open(file_path, "rb") as f:
                raw = f.read()

            # 尝试 UTF-16LE 解码（Word 内部格式）
            text = raw.decode('utf-16-le', errors='ignore')

            # 清理控制字符
            text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', ' ', text)

            # 提取中文文本（连续中文字符，允许中间有少量标点和数字）
            # 使用更宽松的模式：中文字符+任意非ASCII字符的组合
            chinese_pattern = r'[\u4e00-\u9fff][\u4e00-\u9fff\u3000-\u303f\uff00-\uffef\w\s\.\,\;\:\!\?\-\(\)\[\]\d]{3,}'
            matches = re.findall(chinese_pattern, text)

            # 清理和去重
            cleaned = []
            seen = set()
            for m in matches:
                # 移除多余空格
                m = re.sub(r'\s+', ' ', m).strip()
                # 只保留包含至少2个中文字符的文本
                chinese_count = len(re.findall(r'[\u4e00-\u9fff]', m))
                if chinese_count >= 2 and len(m) > 3 and m not in seen:
                    seen.add(m)
                    cleaned.append(m)

            if cleaned:
                return "\n".join(cleaned)

            # 如果上面的方法没提取到足够的文本，尝试更简单的方法
            # 直接提取所有连续的中文字符序列
            chinese_sequences = re.findall(r'[\u4e00-\u9fff]{2,}', text)
            if chinese_sequences:
                # 合并相邻的短序列
                merged = []
                current = ""
                for seq in chinese_sequences:
                    if len(current) + len(seq) < 100:
                        current += seq
                    else:
                        if current:
                            merged.append(current)
                        current = seq
                if current:
                    merged.append(current)
                return "\n".join(merged)

        except Exception as e:
            logger.debug(f"UTF-16LE extraction failed: {e}")

        # 方法4：GBK 解码回退
        try:
            import re
            with open(file_path, "rb") as f:
                raw = f.read()

            for encoding in ["gbk", "gb2312", "gb18030"]:
                try:
                    text = raw.decode(encoding, errors="ignore")
                    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', ' ', text)
                    chinese_sequences = re.findall(r'[\u4e00-\u9fff]{2,}', text)
                    if chinese_sequences and len("".join(chinese_sequences)) > 50:
                        return "\n".join(chinese_sequences)
                except Exception:
                    continue
        except Exception:
            pass

        logger.warning(f"DOC 文本提取失败（无可用提取方式）: {file_path}")
        return ""
    except Exception as e:
        logger.warning(f"DOC 文本提取失败 {file_path}: {e}")
        return ""


def extract_text(file_path: str) -> str:
    """根据文件类型自动提取文本"""
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    elif ext == ".doc":
        return extract_text_from_doc(file_path)
    return ""
